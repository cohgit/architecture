import speech_recognition as sr
import nltk
from nltk.tokenize import sent_tokenize
from docx import Document
from datetime import datetime

# Configurar el reconocimiento de voz
def transcribe_audio(audio_file):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)
    return recognizer.recognize_google(audio, language='es-ES')

# Resumir el texto
def summarize_text(text):
    sentences = sent_tokenize(text)
    # Aquí puedes aplicar un algoritmo de resumen
    summary = sentences[:5]  # Toma las primeras 5 oraciones como ejemplo
    return ' '.join(summary)

# Generar la minuta en Word
def generate_minute(summary, output_file='minuta.docx'):
    document = Document()
    document.add_heading('Minuta de Sesión Técnica', 0)
    document.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
    document.add_heading('Resumen:', level=1)
    document.add_paragraph(summary)
    document.save(output_file)

# Función principal
def main():
    # Transcribir el audio (si tienes un archivo de audio)
    audio_text = transcribe_audio('/home/cogalde/dev/architecture/ImageMaker/BAC/SesionPresencialNuevosPerfiles/Session 1.wav')
    
    # Resumir el texto
    summary = summarize_text(audio_text)

    # Generar la minuta
    generate_minute(summary)

    print("Minuta generada con éxito.")

if __name__ == "__main__":
    main()
